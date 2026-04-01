"""
Generate Vendor Wrap Report Summary - VitalFit Spring Awareness Campaign
Analyzes three vendor wrap reports against the original campaign brief.
Outputs a professionally formatted Word document.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Color scheme ──
NAVY = RGBColor(0x1B, 0x2A, 0x4A)
DARK_BLUE = RGBColor(0x2C, 0x3E, 0x6B)
MED_BLUE = RGBColor(0x3A, 0x5B, 0x9F)
LIGHT_BLUE = RGBColor(0xD6, 0xE4, 0xF0)
RED_FLAG = RGBColor(0xC0, 0x39, 0x2B)
GREEN_WIN = RGBColor(0x1E, 0x7D, 0x32)
AMBER = RGBColor(0xE6, 0x7E, 0x22)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Create a styled table with navy header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)
        run.font.name = "Calibri"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "1B2A4A")

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]

            # Check if this is a specially formatted value
            text = str(cell_val)
            run = p.add_run(text)
            run.font.size = Pt(9)
            run.font.name = "Calibri"
            run.font.color.rgb = DARK_GRAY

            # Highlight miss/fail values in red, wins in green
            if any(kw in text.lower() for kw in ["miss", "fail", "late", "pending", "over budget", "not significant"]):
                run.font.color.rgb = RED_FLAG
                run.bold = True
            elif any(kw in text.lower() for kw in ["beat", "exceeded", "on target", "on time"]):
                run.font.color.rgb = GREEN_WIN
                run.bold = True

            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Zebra striping
            if r_idx % 2 == 1:
                set_cell_shading(cell, "EEF2F7")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def add_section_heading(doc, text, level=2):
    """Add a heading with navy color."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = NAVY
    return heading


def add_bullet(doc, text, bold_prefix=None, color=None):
    """Add a bullet point, optionally with a bold prefix."""
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        if color:
            run.font.color.rgb = color
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        run.font.color.rgb = DARK_GRAY
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        if color:
            run.font.color.rgb = color
        else:
            run.font.color.rgb = DARK_GRAY
    return p


def add_callout(doc, text, color=RED_FLAG):
    """Add a bold callout paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    run.font.color.rgb = color
    return p


def build_document():
    doc = Document()

    # ── Set default font ──
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    font.color.rgb = DARK_GRAY

    # ── Adjust heading styles ──
    for level in range(1, 4):
        h_style = doc.styles[f"Heading {level}"]
        h_style.font.name = "Calibri"
        h_style.font.color.rgb = NAVY

    # ══════════════════════════════════════════════════════════════════
    # TITLE PAGE
    # ══════════════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("VitalFit Spring Awareness Campaign")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = NAVY
    run.font.name = "Calibri"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Vendor Wrap Report Analysis")
    run.font.size = Pt(18)
    run.font.color.rgb = MED_BLUE
    run.font.name = "Calibri"

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("March 2026 Flight  |  Prepared April 2026")
    run.font.size = Pt(12)
    run.font.color.rgb = DARK_GRAY
    run.font.name = "Calibri"

    doc.add_paragraph()
    vendors_line = doc.add_paragraph()
    vendors_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = vendors_line.add_run("Vendors Reviewed: ReachMax Digital  |  TikTok  |  SparkPoint Media")
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_BLUE
    run.font.name = "Calibri"

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # EXECUTIVE OVERVIEW
    # ══════════════════════════════════════════════════════════════════
    add_section_heading(doc, "Executive Overview", level=1)

    p = doc.add_paragraph()
    run = p.add_run(
        "This document reviews the wrap reports submitted by three vendors for the VitalFit Spring "
        "Awareness campaign (March 1-31, 2026). Each vendor report is assessed against the original "
        "campaign brief and SOW terms. The analysis identifies where vendors genuinely delivered, "
        "where they are spinning underperformance as success, and where data is missing or suspect. "
        "Total campaign budget across all three vendors was $110,000."
    )
    run.font.size = Pt(10)
    run.font.name = "Calibri"

    doc.add_paragraph()

    add_styled_table(doc,
        ["Vendor", "Channel", "Budget", "Spend", "Variance", "Report Timeliness"],
        [
            ["ReachMax Digital", "Programmatic Display", "$25,000", "$24,800", "-$200 (under)", "2 days late"],
            ["TikTok", "Creator / Spark Ads", "$55,000", "$58,870", "+$3,870 (over budget)", "4 days late"],
            ["SparkPoint Media", "Connected TV / OTT", "$30,000", "$30,200", "+$200 (on target)", "On time"],
        ],
        col_widths=[3.5, 3.5, 2.2, 2.2, 3.0, 2.5]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # VENDOR 1: REACHMAX DIGITAL
    # ══════════════════════════════════════════════════════════════════
    add_section_heading(doc, "Vendor 1: ReachMax Digital -- Programmatic Display", level=1)

    # ── 1.1 Results vs Objectives ──
    add_section_heading(doc, "Headline Results vs Objectives", level=2)

    add_styled_table(doc,
        ["Metric", "Target", "Delivered", "Verdict"],
        [
            ["Total Impressions", "3,000,000", "3,420,000 (114%)", "Beat target"],
            ["Viewability", "80%+", "70%", "MISS -- 10pts below target"],
            ["Viewable Impressions", "2,400,000", "2,394,000 (100%)", "Barely met -- inflated by overdelivery"],
            ["CTR", "0.12%", "0.12%", "On target -- not exceeded"],
            ["Total Clicks", "3,600", "4,104 (114%)", "Beat -- driven by impression volume"],
            ["CPM (Total)", "$8.33", "$7.25", "Beat target"],
            ["CPM (Viewable)", "N/A", "$10.36", "True cost 43% higher than headline CPM"],
            ["Ontario Geo", "40%", "38%", "Slight miss"],
            ["Quebec Geo", "25%", "18%", "MISS -- significant, 7pts below"],
            ["Brand Lift Study", "Primary KPI", "Pending (delayed)", "FAIL -- no data delivered"],
            ["Report Delivery", "10 biz days", "12 biz days", "Late -- 2 days past SOW terms"],
        ],
        col_widths=[3.5, 2.5, 3.5, 4.5]
    )

    doc.add_paragraph()

    # ── 1.2 Vendor Spin Detection ──
    add_section_heading(doc, "Vendor Spin Detection", level=2)

    add_bullet(doc, ' "Over-delivered on impressions by 14%" -- True on raw impressions, but viewability was only 70% vs 80% target. The overdelivery was achieved by buying cheaper, low-viewability inventory. Viewable impressions (the actual primary KPI) barely reached 100% of target. The vendor is celebrating a vanity metric to distract from the real underperformance.', "SPIN:", RED_FLAG)
    add_bullet(doc, ' "CPM efficiency beat target by 13%" -- Headline CPM of $7.25 looks great, but viewable CPM is $10.36. They bought cheap inventory that humans did not actually see. The real cost per viewable impression is 43% higher than the headline figure they are promoting.', "SPIN:", RED_FLAG)
    add_bullet(doc, ' "Strong engagement with fitness-oriented audiences" -- CTR was exactly 0.12%, matching the benchmark but not exceeding it. This is par, not strong. One subsegment (Health & Fitness Enthusiasts) hit 0.18% but overall performance was average.', "SPIN:", RED_FLAG)
    add_bullet(doc, ' Geographic shortfall buried in a table -- Quebec delivered 18% vs 25% target, a 28% shortfall. The report makes no mention of this miss in the executive summary or key wins. Ontario also slightly underdelivered at 38% vs 40%.', "SPIN:", RED_FLAG)

    doc.add_paragraph()

    # ── 1.3 Genuine Wins ──
    add_section_heading(doc, "Genuine Wins", level=2)

    add_bullet(doc, "Came in under budget ($24,800 vs $25,000) -- no overspend.", color=GREEN_WIN)
    add_bullet(doc, "Health & Fitness Enthusiasts segment at 0.18% CTR -- 50% above benchmark. Worth retargeting.", color=GREEN_WIN)
    add_bullet(doc, "68% of impressions reached core demo (Adults 25-44) -- reasonable audience alignment.", color=GREEN_WIN)
    add_bullet(doc, "Weekly pacing reports delivered on schedule throughout the flight.", color=GREEN_WIN)

    doc.add_paragraph()

    # ── 1.4 Red Flags ──
    add_section_heading(doc, "Red Flags", level=2)

    add_bullet(doc, ' Brand lift study "pending" -- this is a PRIMARY KPI per the SOW. Campaign ended March 31, report delivered April 14, and the lift study still has no results. The survey partner being "delayed" is the vendor\'s responsibility to manage. No data = the most important KPI cannot be evaluated.', "CRITICAL: ", RED_FLAG)
    add_bullet(doc, ' "Zero brand safety violations" claim has no supporting methodology. No third-party verification report attached. No list of excluded domains or categories. This is an unsubstantiated claim.', "UNVERIFIED: ", RED_FLAG)
    add_bullet(doc, ' Report delivered 2 days late per SOW terms (12 business days vs 10). Not flagged or acknowledged by the vendor.', "LATE: ", AMBER)
    add_bullet(doc, ' Placement-level data "not included in this summary" -- 42 rows referenced but not shared. This data is needed to verify brand safety and viewability claims.', "MISSING DATA: ", AMBER)

    doc.add_paragraph()

    # ── 1.5 Key Takeaways ──
    add_section_heading(doc, "Key Takeaways for Campaign Deck", level=2)

    add_bullet(doc, "Programmatic display delivered volume but not quality. Viewability miss means ~600K impressions were served but never actually seen by a human.")
    add_bullet(doc, "Brand lift data -- the primary reason this channel was in the plan -- is not available. Cannot assess awareness impact.")
    add_bullet(doc, "Quebec underdelivery (18% vs 25%) is a pattern that may repeat across vendors. Investigate French-language inventory availability.")
    add_bullet(doc, "For Q2: require viewable CPM as the primary efficiency metric in the SOW, not total CPM. Add contractual penalties for late reporting and missing KPI data.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # VENDOR 2: TIKTOK
    # ══════════════════════════════════════════════════════════════════
    add_section_heading(doc, "Vendor 2: TikTok -- Creator / Spark Ads", level=1)

    # ── 2.1 Results vs Objectives ──
    add_section_heading(doc, "Headline Results vs Objectives", level=2)

    add_styled_table(doc,
        ["Metric", "Target", "Delivered", "Verdict"],
        [
            ["Impressions", "5,000,000", "5,290,000 (106%)", "Beat target"],
            ["Video Views (2s+)", "3,500,000", "3,808,000 (109%)", "Beat target"],
            ["Video Views (6s+)", "2,000,000", "2,380,000 (119%)", "Beat target"],
            ["Clicks", "50,000", "62,340 (125%)", "Beat target"],
            ["CTR", "1.00%", "1.18%", "Beat target"],
            ["CPC", "$1.10", "$0.94", "Beat target"],
            ["CPM", "$11.00", "$11.13", "On target (slight overage)"],
            ["Total Spend", "$55,000", "$58,870 (107%)", "OVER BUDGET by $3,870"],
            ["Quebec Engagement", "Comparable", "40% lower", "MISS -- French content gap"],
            ["Report Delivery", "10 biz days", "14 biz days", "Late -- 4 days past SOW terms"],
        ],
        col_widths=[3.5, 2.5, 3.5, 4.5]
    )

    doc.add_paragraph()

    # ── 2.2 Vendor Spin Detection ──
    add_section_heading(doc, "Vendor Spin Detection", level=2)

    add_bullet(doc, ' Budget overspend framed as positive: "Spend came in 7% over budget due to strong performance in final week -- auto-scaling was enabled." The vendor spent $3,870 more than authorized without approval. Framing unauthorized overspend as a feature of "strong performance" is textbook spin. Auto-scaling should have had a budget cap.', "SPIN:", RED_FLAG)
    add_bullet(doc, ' Purchase attribution uses 28-day click window: 412 purchases at $142.89 CPA sounds reasonable, but a 28-day click attribution window is extremely generous. Industry standard for social is 7-day click / 1-day view. This inflates conversion counts significantly. The true CPA on a standard window is likely 2-3x higher.', "SPIN:", RED_FLAG)
    add_bullet(doc, ' "One of the top-performing supplement campaigns on TikTok Canada in Q1 2026" -- This is an unverifiable claim. TikTok does not publish campaign-level benchmarks publicly. This is vendor puffery with no data to support it.', "SPIN:", RED_FLAG)
    add_bullet(doc, ' Recommendation to "increase budget to $80K/month" -- a 45% budget increase recommendation from a vendor who already overspent by 7% without approval. This is a classic upsell disguised as strategic advice.', "SPIN:", AMBER)

    doc.add_paragraph()

    # ── 2.3 Genuine Wins ──
    add_section_heading(doc, "Genuine Wins", level=2)

    add_bullet(doc, "CTR of 1.18% genuinely beat the 1.00% target -- 18% above benchmark. Creator content drove real engagement.", color=GREEN_WIN)
    add_bullet(doc, "CPC of $0.94 vs $1.10 target -- 15% more efficient. Legitimate efficiency gain.", color=GREEN_WIN)
    add_bullet(doc, "@NutritionNerd (180K followers) had the highest engagement rate (5.6%) and lowest CPC ($0.68). Micro-creator outperformance is a real insight worth acting on.", color=GREEN_WIN)
    add_bullet(doc, "72% of conversions from Spark Ads vs 28% standard in-feed. Creator-driven formats genuinely outperformed brand content.", color=GREEN_WIN)
    add_bullet(doc, "6-second view rate strong at 119% of target -- content held attention beyond the scroll.", color=GREEN_WIN)

    doc.add_paragraph()

    # ── 2.4 Red Flags ──
    add_section_heading(doc, "Red Flags", level=2)

    add_bullet(doc, " Unauthorized overspend of $3,870 (7%). Auto-scaling without a hard budget cap is a vendor management failure. Require pre-approval for any spend above 100% of budget.", "CRITICAL:", RED_FLAG)
    add_bullet(doc, " @GymBro_Tyler: highest follower count (1.2M) but lowest engagement rate (2.1%) and highest CPC ($1.18). Spent the most for the least efficient results. Vendor acknowledges this but still ran 3 videos with this creator.", "INEFFICIENCY:", AMBER)
    add_bullet(doc, " 28-day click attribution inflates all lower-funnel numbers. Request 7-day click / 1-day view numbers for apples-to-apples comparison.", "ATTRIBUTION:", RED_FLAG)
    add_bullet(doc, " Report was 4 days late (14 business days vs 10 SOW requirement). Worst of the three vendors on timeliness.", "LATE:", AMBER)
    add_bullet(doc, " Quebec engagement 40% below other provinces. French-language creator content gap acknowledged but no plan was in place before launch.", "GEO MISS:", AMBER)

    doc.add_paragraph()

    # ── 2.5 Key Takeaways ──
    add_section_heading(doc, "Key Takeaways for Campaign Deck", level=2)

    add_bullet(doc, "TikTok delivered strong upper-funnel engagement. CTR, CPC, and view metrics all above target. The channel works for awareness and consideration.")
    add_bullet(doc, "Lower-funnel performance claims are inflated by a 28-day attribution window. Do not cite purchase CPA ($142.89) without this caveat. Request 7-day data before including in performance deck.")
    add_bullet(doc, "Micro-creators (under 500K followers) outperformed the large creator across all efficiency metrics. Reallocate creator budget away from reach-focused influencers toward engagement-focused ones.")
    add_bullet(doc, "For Q2: enforce hard budget caps, require pre-approval for auto-scaling, mandate standard attribution windows in the SOW, and invest in French-language creator partnerships for Quebec.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # VENDOR 3: SPARKPOINT MEDIA
    # ══════════════════════════════════════════════════════════════════
    add_section_heading(doc, "Vendor 3: SparkPoint Media -- Connected TV / OTT", level=1)

    # ── 3.1 Results vs Objectives ──
    add_section_heading(doc, "Headline Results vs Objectives", level=2)

    add_styled_table(doc,
        ["Metric", "Target", "Delivered", "Verdict"],
        [
            ["Total Impressions", "2,000,000", "2,100,000 (105%)", "Beat target"],
            ["Completed Views (VCR)", "1,600,000 (80%)", "1,722,000 (82%)", "Exceeded target"],
            ["CPM", "$15.00", "$14.38", "Beat target"],
            ["CPCV", "$0.019", "$0.018", "Beat target"],
            ["Total Spend", "$30,000", "$30,200 (101%)", "On target"],
            ["Unique Households", "N/A", "680,000", "No target set -- baseline data"],
            ["Quebec Geo", "25% (implied)", "19%", "MISS -- below target"],
            ["Brand Lift: Awareness", "Lift target", "+6 pts (95% CI)", "On target"],
            ["Brand Lift: Ad Recall", "Lift target", "+16 pts (99% CI)", "Beat target"],
            ["Brand Lift: Purchase Intent", "Lift target", "+3 pts (88% CI)", "Not statistically significant"],
            ["Report Delivery", "10 biz days", "9 biz days", "On time"],
        ],
        col_widths=[3.5, 2.8, 3.5, 4.2]
    )

    doc.add_paragraph()

    # ── 3.2 Vendor Spin Detection ──
    add_section_heading(doc, "Vendor Spin Detection", level=2)

    add_bullet(doc, ' "Strong brand lift across all measured metrics" -- Purchase intent lift was +3 points at only 88% confidence. This is NOT statistically significant by standard thresholds (95%). Brand favorability at 91% confidence is also below the accepted threshold. Only awareness (+6pts, 95%) and ad recall (+16pts, 99%) are statistically valid. Claiming "all metrics" showed strong lift is misleading.', "SPIN:", RED_FLAG)
    add_bullet(doc, ' Recommendations to "increase budget to $50K/month" and add Hulu + Disney+ -- SparkPoint does not discuss the budget implications of these recommendations. Adding premium streaming inventory could double the CPM. The recommendation is a vendor upsell presented as strategy.', "SPIN:", AMBER)
    add_bullet(doc, ' "Efficient CPM vs CTV category benchmark of $18-22" -- The benchmark cited is suspiciously high. SparkPoint ran primarily on Roku Channel, Tubi, Pluto TV, and Samsung TV+ -- all free ad-supported streaming (FAST) platforms. The $18-22 benchmark is for premium CTV (Hulu, Disney+). Comparing FAST platform CPMs to premium benchmarks flatters their efficiency.', "SPIN:", AMBER)

    doc.add_paragraph()

    # ── 3.3 Genuine Wins ──
    add_section_heading(doc, "Genuine Wins", level=2)

    add_bullet(doc, "This is the strongest-performing vendor of the three. Delivered on nearly all core metrics -- impressions, VCR, CPM, CPCV.", color=GREEN_WIN)
    add_bullet(doc, "Ad recall lift of +16 points at 99% confidence is a genuinely strong result. CTV drove meaningful brand awareness.", color=GREEN_WIN)
    add_bullet(doc, "72% primetime delivery (7-11pm) -- ads were seen during engaged viewing occasions, not 3am backfill.", color=GREEN_WIN)
    add_bullet(doc, "Report delivered on time (9 business days). The only vendor to meet the SOW deadline.", color=GREEN_WIN)
    add_bullet(doc, "Average frequency of 3.1x per household aligns well with awareness objectives without over-saturating.", color=GREEN_WIN)
    add_bullet(doc, "Lucid brand lift methodology attached (3-page PDF) -- transparent about measurement approach, unlike other vendors.", color=GREEN_WIN)

    doc.add_paragraph()

    # ── 3.4 Red Flags ──
    add_section_heading(doc, "Red Flags", level=2)

    add_bullet(doc, " Purchase intent lift (+3pts, 88% CI) is not statistically significant. Do not include in client reporting as a proven result.", "STAT SIGNIFICANCE:", RED_FLAG)
    add_bullet(doc, " Quebec at 19% -- below 25% target. Same French-language inventory issue seen across all three vendors. SparkPoint acknowledges 'French-language inventory limited on Tubi/Pluto' but did not proactively address this during flight.", "GEO MISS:", AMBER)
    add_bullet(doc, " Inventory skews heavily toward FAST platforms (Roku, Tubi, Pluto, Samsung TV+). These are legitimate CTV but are the lower end of the quality spectrum. If VitalFit wants premium brand association, a shift toward paid streaming services would be needed at higher cost.", "INVENTORY QUALITY:", AMBER)
    add_bullet(doc, " No frequency capping data shared. Average of 3.1x is fine but distribution could be skewed -- some users may have seen the ad 10+ times while others saw it once.", "MISSING DATA:", AMBER)

    doc.add_paragraph()

    # ── 3.5 Key Takeaways ──
    add_section_heading(doc, "Key Takeaways for Campaign Deck", level=2)

    add_bullet(doc, "CTV was the most reliable channel in the mix. SparkPoint delivered on time, on budget, and on target for core metrics. This should be the anchor of Q2 planning.")
    add_bullet(doc, "Brand lift results are mixed: ad recall is strong and significant, but purchase intent lift is not statistically proven. Present awareness and recall results confidently; flag purchase intent as inconclusive.")
    add_bullet(doc, "Consider a hybrid CTV strategy for Q2: maintain FAST platforms for efficient reach, add a smaller premium tier (Hulu/Disney+) for brand halo, and set distinct KPIs for each tier.")
    add_bullet(doc, "Quebec underdelivery is a systemic issue across all vendors. Consider a dedicated French-language media buy or Quebec-specific partner for Q2.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # CROSS-VENDOR SUMMARY
    # ══════════════════════════════════════════════════════════════════
    add_section_heading(doc, "Cross-Vendor Summary", level=1)

    # ── Overall Performance ──
    add_section_heading(doc, "Overall Campaign Performance", level=2)

    p = doc.add_paragraph()
    run = p.add_run(
        "The VitalFit Spring Awareness campaign ran across three vendors with a combined budget of $110,000. "
        "Total spend was $113,870 -- 3.5% over budget, driven entirely by TikTok's unauthorized auto-scaling overspend. "
        "The campaign achieved its volume targets across all channels but fell short on quality metrics (viewability, "
        "geo distribution) and measurement deliverables (brand lift data delayed/incomplete). "
        "The Quebec underperformance (18-19% vs 25% target across all three vendors) represents a systemic French-language "
        "inventory and content gap that no vendor adequately addressed."
    )
    run.font.size = Pt(10)
    run.font.name = "Calibri"

    doc.add_paragraph()

    add_styled_table(doc,
        ["Dimension", "ReachMax Digital", "TikTok", "SparkPoint Media"],
        [
            ["Budget Management", "Under budget (-$200)", "Over budget (+$3,870)", "On budget (+$200)"],
            ["Core KPI Delivery", "Mixed -- volume yes, quality no", "Strong on engagement metrics", "Strong across the board"],
            ["Quebec Target (25%)", "MISS (18%)", "MISS (40% lower engagement)", "MISS (19%)"],
            ["Brand Lift Data", "FAIL -- pending/delayed", "N/A (not in scope)", "Partial -- 2 of 4 significant"],
            ["Report Timeliness", "Late (2 days)", "Late (4 days)", "On time"],
            ["Vendor Transparency", "Low -- missing data, unverified claims", "Medium -- acknowledges issues", "High -- methodology shared"],
            ["Overall Grade", "C-", "B", "A-"],
        ],
        col_widths=[3.5, 3.5, 3.5, 3.5]
    )

    doc.add_paragraph()

    # ── Who Delivered vs Who Fell Short ──
    add_section_heading(doc, "Who Delivered vs Who Fell Short", level=2)

    add_callout(doc, "SparkPoint Media (CTV) -- Best Performer", GREEN_WIN)
    add_bullet(doc, "Only vendor to deliver on time, on budget, and on target across core metrics. Transparent reporting with methodology attached. Brand lift results are partially proven. Earned the right to increased investment in Q2.")

    add_callout(doc, "TikTok -- Strong Engagement, Poor Governance", AMBER)
    add_bullet(doc, "Genuinely strong upper-funnel performance on CTR, CPC, and views. Creator strategy produced real insights (micro-creator outperformance). However, unauthorized overspend, inflated attribution windows, and the latest report delivery undermine trust. Needs tighter controls, not more budget.")

    add_callout(doc, "ReachMax Digital -- Underperformed on What Matters", RED_FLAG)
    add_bullet(doc, "Delivered volume (impressions) but missed on quality (viewability), geography (Quebec), and the primary KPI (brand lift data still pending). Report was late. Brand safety claims are unverified. The vendor's narrative does not match the data. Requires a serious performance conversation before Q2 renewal.")

    doc.add_paragraph()

    # ── Q2 Recommendations ──
    add_section_heading(doc, "Recommendations for Q2 Vendor Strategy", level=2)

    p = doc.add_paragraph()
    run = p.add_run("1. Tighten SOW terms across all vendors.")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    run.font.color.rgb = NAVY
    p2 = doc.add_paragraph()
    run2 = p2.add_run(
        "Add contractual requirements: viewable CPM as primary efficiency metric (not total CPM), "
        "hard budget caps with pre-approval for overspend, standardized attribution windows (7-day click / 1-day view), "
        "penalties for late reporting, and mandatory third-party brand safety verification reports."
    )
    run2.font.size = Pt(10)
    run2.font.name = "Calibri"

    p = doc.add_paragraph()
    run = p.add_run("2. Address Quebec systematically.")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    run.font.color.rgb = NAVY
    p2 = doc.add_paragraph()
    run2 = p2.add_run(
        "All three vendors missed Quebec targets. This is not a vendor problem -- it is a market structure problem. "
        "French-language digital inventory is limited on major platforms. Consider a dedicated Quebec media partner, "
        "French-language creator partnerships for TikTok, and French-language CTV inventory pre-buys for Q2."
    )
    run2.font.size = Pt(10)
    run2.font.name = "Calibri"

    p = doc.add_paragraph()
    run = p.add_run("3. Reallocate budget toward proven performers.")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    run.font.color.rgb = NAVY
    p2 = doc.add_paragraph()
    run2 = p2.add_run(
        "SparkPoint (CTV) earned a budget increase based on performance. TikTok should maintain budget but with "
        "tighter controls. ReachMax should be put on notice: require a performance improvement plan addressing "
        "viewability, brand lift delivery timelines, and geographic targeting before committing Q2 budget."
    )
    run2.font.size = Pt(10)
    run2.font.name = "Calibri"

    p = doc.add_paragraph()
    run = p.add_run("4. Require brand lift study timelines in the SOW.")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    run.font.color.rgb = NAVY
    p2 = doc.add_paragraph()
    run2 = p2.add_run(
        "ReachMax's pending brand lift study -- the primary KPI -- is unacceptable. For Q2, require survey partner "
        "coordination before campaign launch, mid-flight preliminary results, and final data delivered with the wrap report. "
        "If a vendor cannot commit to this timeline, the KPI should be removed from their scope."
    )
    run2.font.size = Pt(10)
    run2.font.name = "Calibri"

    p = doc.add_paragraph()
    run = p.add_run("5. Standardize vendor report cards.")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    run.font.color.rgb = NAVY
    p2 = doc.add_paragraph()
    run2 = p2.add_run(
        "Create a standard template that all vendors must follow for wrap reports. This ensures consistent data, "
        "makes cross-vendor comparison easier, and prevents vendors from burying misses in narrative text. "
        "Include a mandatory section where vendors self-identify what did not meet targets and why."
    )
    run2.font.size = Pt(10)
    run2.font.name = "Calibri"

    # ── Save ──
    output_path = os.path.join(os.path.dirname(__file__), "output.docx")
    doc.save(output_path)
    print(f"Document saved to: {output_path}")


if __name__ == "__main__":
    build_document()
